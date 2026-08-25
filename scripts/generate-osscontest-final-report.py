#!/usr/bin/env python3
from __future__ import annotations

import json
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "artifacts/submission/2026-osscontest"
TEMPLATE = BASE / "template/official-result-report-template.docx"
OUT = BASE / "final"
SBOM_REVIEW = BASE / "sbom/license-review.json"
RECEIPT = os.environ.get("OSSCONTEST_RECEIPT_NUMBER", "접수번호")
TEAM = os.environ.get("OSSCONTEST_TEAM_NAME", "ClaimGate")
DOCX = OUT / f"2026 오픈소스 개발자대회 결과보고서_{RECEIPT}({TEAM}).docx"
PDF = OUT / f"2026 오픈소스 개발자대회 결과보고서_{RECEIPT}({TEAM}).pdf"
REPO_URL = "https://github.com/WooYoungSang/warvis-claimgate"
VIDEO_URL = os.environ.get("OSSCONTEST_VIDEO_URL", "[YOUTUBE_URL_REQUIRED]")


def existing_paths(paths: list[str]) -> str:
    existing = [path for path in paths if (ROOT / path).exists()]
    return ", ".join(existing) if existing else "공개 관리 문서 생성 전"


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold


def add_lines(cell, lines: list[str]) -> None:
    cell.text = ""
    for index, text in enumerate(lines):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.style = cell.paragraphs[0].style
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)


def set_all_font(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Malgun Gothic"
            run.font.size = Pt(10)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Malgun Gothic"
                        run.font.size = Pt(10)


def add_picture(cell, path: Path, width: float) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    if not TEMPLATE.exists():
        raise SystemExit(f"missing official template: {TEMPLATE}")
    doc = Document(TEMPLATE)
    tables = doc.tables

    # 참가 정보
    info = tables[2]
    set_cell(info.cell(1, 1), TEAM)
    set_cell(info.cell(1, 3), "1명")
    set_cell(info.cell(2, 1), "일반")
    set_cell(info.cell(2, 3), "자유과제")

    report = tables[3]
    set_cell(report.cell(1, 1), "ClaimGate")
    set_cell(report.cell(2, 1), REPO_URL)
    set_cell(report.cell(3, 1), VIDEO_URL)
    set_cell(report.cell(4, 1), "AI·공공데이터 출력의 주장을 Source Anchor, 결정론적 위험 규칙, 사람의 검토, Evidence Pack으로 연결하는 오픈소스 claim review framework")
    add_lines(report.cell(6, 1), [
        "• 생성형 AI 문장은 자연스러워도 출처·값·날짜·기관이 틀릴 수 있다. ClaimGate는 생성 이후의 검토 책임과 재사용 조건을 코드로 명시한다.",
        "• 목표: AI는 후보 주장·근거만 제안하고, Source Anchor·명시적 규칙·사람의 최종 판정으로 승인된 주장만 Evidence Pack에 넣는다.",
        "• 불변식: No Anchor, No Claim / AI Curator, Not Judge / deterministic risk / Evidence Pack First.",
    ])
    add_lines(report.cell(7, 1), [
        "• Node.js 20+, TypeScript strict, pnpm monorepo, React 18, Vite, Vitest, tsup",
        "• 선택적 로컬 AI: Gemma 4 12B candidate extractor, repo-local sparse RAG, bounded QLoRA adapter evaluation",
        "• Go 1.22 도구: kbctl JSON 지식정본 CLI, FMON Bubble Tea read-model TUI",
        "• 기본 경로는 서버·DB·인증·호스팅 LLM 없이 offline deterministic fixture-first로 재현 가능",
    ])
    add_lines(report.cell(8, 1), [
        "• @claimgate/core: Claim·Source Anchor·review state·risk trace·Evidence Pack 불변식",
        "• @claimgate/ui: 숨은 판정 권한이 없는 controlled React components",
        "• @claimgate/pack-*: civic / health / MOFA ODA fixture와 도메인 판단 규칙",
        "• @claimgate/ai-local: 후보 추출과 RAG provenance만 담당하며 검증·위험결정·투영 권한 없음",
        "• 데이터 흐름: Candidate → accepted Source Anchor → deterministic risk → Human Review → Evidence Pack → Report/Graph",
    ])
    add_picture(report.cell(8, 1), ROOT / "artifacts/submission/2026-mofa-ai/final-presentation/assets/01-guided-demo-start.png", 5.4)
    add_lines(report.cell(9, 1), [
        "[서면 1/5 프로젝트 구조 및 코드 완성도] strict TypeScript 패키지 경계와 순수 core를 유지한다. 출처 없는 verified/corrected 전이, AI의 위험·판정·투영, 미승인 Evidence Pack을 fail-closed로 거부한다.",
        "[서면 2/5 오픈소스 프로젝트로의 발전 가능성] core/UI 계약을 고정하고 DomainPack이 fixture·rule·copy를 확장한다. civic·health·MOFA ODA 3개 팩과 conformance kit가 새 팩의 호환성을 검사한다.",
        "[서면 3/5 개발 문서의 구체성] README의 Node 20+/pnpm quickstart, CONTRIBUTING의 검증 절차, SECURITY의 신뢰경계, 패키지별 API 문서로 설치→데모→확장을 재현한다.",
        "[실행·기능테스트] pnpm install --frozen-lockfile → pnpm demo → pnpm eval:framework. unit/conformance/e2e/runbook/clean-clone/deployment/submission negative controls를 workspace에서 자동 실행한다.",
    ])
    add_picture(report.cell(9, 1), ROOT / "artifacts/submission/2026-mofa-ai/final-presentation/assets/02-review-workspace.png", 5.4)
    add_lines(report.cell(10, 1), [
        "[활용성] 공공기관·연구·보건·시민데이터의 AI 초안을 주장 단위 reviewer workflow로 바꾸고, 근거 위치·규칙 trace·사람의 결정을 함께 인계한다.",
        "[작품 데모(완성도)] offline fixture만으로 RED/YELLOW/GREEN 분류, 정정·검증, Evidence Pack export가 반복 실행된다. local Gemma/RAG/QLoRA는 선택 기능이며 core 판정 권한을 갖지 않는다.",
        "[커뮤니티 확장 가능성] CONTRIBUTING·Issue 양식·conformance kit로 DomainPack 제안과 결함 보고 경로를 제공한다. 외부 contributor·PR 채택 실적은 아직 없으며 향후 공개 운영으로 검증한다.",
        "[향후 검증] 읽기 전용 공공데이터 adapter와 제한 파일럿으로 검토시간·재확인시간·정정 재사용률을 측정한다. 현재 운영 정확도·고객 채택·시간 절감 수치는 주장하지 않는다.",
    ])
    add_lines(report.cell(11, 1), [
        "[서면 4/5 프로젝트 혁신성] 더 많은 답을 생성하는 대신, 주장 단위 Source Anchor→deterministic risk→Human Review→Evidence Pack을 하나의 검토 프로토콜로 구현한다.",
        "[서면 5/5 프로젝트 협업 및 관리체계] MIT LICENSE, CONTRIBUTING, SECURITY, kbctl SSOT, 120+ commits와 자동 lint/typecheck/test를 운영한다. 현재 1인 개발이므로 외부 Review·PR·커뮤니티 활동 이력을 과장하지 않는다.",
        "[오픈소스SW 적절성·라이선스 검증] React/Vite/Vitest/Go OSS를 정상 실행 경로에 사용하고, LICENSE·THIRD_PARTY_LICENSES·붙임 SBOM·AI 모델 명세로 식별·충돌 여부를 점검한다.",
        "[관리 증거] " + existing_paths(["CODE_OF_CONDUCT.md", "GOVERNANCE.md", "ROADMAP.md", ".github/ISSUE_TEMPLATE/bug_report.yml", ".github/ISSUE_TEMPLATE/domain_pack.yml"]),
        "[한계·로드맵] live OpenAPI·production vector DB·실사용 효과는 future scope다. framework 안정화→ODA 시제품→파일럿→community DomainPack registry 순으로 공개 검증한다.",
    ])

    if not SBOM_REVIEW.is_file():
        raise SystemExit(f"missing generated license review: {SBOM_REVIEW}")
    review = json.loads(SBOM_REVIEW.read_text(encoding="utf-8"))
    if review.get("verdict") != "PASS" or review.get("unresolvedRisks") != []:
        raise SystemExit("license review is not submission-ready")
    sbom = [
        (
            component["name"],
            component["version"],
            component["license"],
            component["repository"],
            f"{component['scope']}: {component['purpose']}",
        )
        for component in review["components"]
    ]
    sbom_table = tables[5]
    for row in sbom_table.rows[1:]:
        for cell in row.cells:
            set_cell(cell, "")
    while len(sbom_table.rows) < len(sbom) + 1:
        sbom_table.add_row()
    for index, item in enumerate(sbom, 1):
        values = (str(index),) + item
        for col, value in enumerate(values):
            set_cell(sbom_table.cell(index, col), value)

    ai = tables[8]
    set_cell(ai.cell(1, 0), "□ 유형 1\n▣ 유형 2: Gemma 4 12B 기반 candidate-only QLoRA 추가학습(로컬 선택 경로). 기본 deterministic fixture 경로는 모델 없이 동작.\n□ 유형 3")
    set_cell(ai.cell(3, 1), "Gemma 4 12B (Google), Ollama model tag gemma4:12b")
    set_cell(ai.cell(3, 3), "Apache-2.0 — Gemma 4 전용 라이선스 적용")
    set_cell(ai.cell(5, 1), "MOFA ODA 공개 demo fixture 유래 9건: train 6 / fixture family별 non-overlap bounded holdout 3. 고객·비공개 데이터 없음.")
    set_cell(ai.cell(6, 1), "CandidateClaim 허용 필드만 유지하고 train/holdout을 분리. response-only loss와 EOS로 strict JSON completion을 학습.")
    set_cell(ai.cell(7, 1), "미배포 — 모델 및 adapter weights는 공개 저장소와 제출 ZIP에 포함하지 않음. 재현 스크립트와 bounded 평가 JSON만 대표 저장소에 공개.")
    set_cell(ai.cell(8, 1), "Gemma 4 12B QLoRA 60-step, local 4-bit 학습. holdout strict JSON/exact text 3/3, authority violation 0. productionQuality=false이며 운영 정확도 주장이 아님.")
    set_cell(ai.cell(10, 1), "MIT License")
    set_cell(ai.cell(10, 3), REPO_URL)
    set_cell(ai.cell(11, 1), "OpenAI Codex 계열 도구를 코드 작성·문서화·리뷰 보조에 활용. 모든 변경은 테스트·typecheck·lint·negative controls로 검증했으며 상용 AI API는 제품 runtime 필수 의존성이 아님.")

    # 제출 시 삭제하도록 명시된 첫 안내 표 제거.
    guide = doc.tables[0]._element
    guide.getparent().remove(guide)
    set_all_font(doc)
    doc.save(DOCX)

    pdf_dir = Path("/tmp/claimgate-osscontest-report-pdf")
    pdf_dir.mkdir(exist_ok=True)
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(DOCX)
    ], check=True)
    generated = pdf_dir / (DOCX.stem + ".pdf")
    PDF.write_bytes(generated.read_bytes())
    print(DOCX)
    print(PDF)


if __name__ == "__main__":
    main()
