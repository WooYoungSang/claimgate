#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import json
import re
import subprocess
import urllib.error
import urllib.request
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts/submission/2026-osscontest/final"


def require(value: bool, message: str) -> None:
    if not value:
        raise SystemExit(f"osscontest-final: FAIL: {message}")


def digest(path: Path) -> str:
    value = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            value.update(chunk)
    return value.hexdigest()


def pdf_pages(path: Path) -> list[str]:
    info = subprocess.check_output(["pdfinfo", str(path)], text=True)
    match = re.search(r"^Pages:\s+(\d+)$", info, re.MULTILINE)
    require(match is not None, "could not determine PDF page count")
    count = int(match.group(1))
    return [
        subprocess.check_output(
            ["pdftotext", "-f", str(page), "-l", str(page), str(path), "-"],
            text=True,
        )
        for page in range(1, count + 1)
    ]


def public_repository_blocker(url: str) -> str | None:
    if not re.fullmatch(r"https://github\.com/[^/\s]+/[^/\s]+/?", url):
        return "PUBLIC_REPOSITORY_URL_INVALID"
    request = urllib.request.Request(url, headers={"User-Agent": "ClaimGate submission validator"})
    try:
        with urllib.request.urlopen(request, timeout=10) as response:
            if response.status != 200:
                return f"PUBLIC_REPOSITORY_HTTP_{response.status}"
    except urllib.error.HTTPError as error:
        return f"PUBLIC_REPOSITORY_HTTP_{error.code}"
    except (urllib.error.URLError, TimeoutError):
        return "PUBLIC_REPOSITORY_UNVERIFIED"
    return None


def main() -> None:
    manifest = json.loads((OUT / "SUBMISSION-MANIFEST.json").read_text())
    docx = OUT / manifest["files"][0]["path"]
    pdf = OUT / manifest["files"][1]["path"]
    video = OUT / "claimgate-osscontest-demo.mp4"
    package = OUT / manifest["uploadZip"]
    for path in [docx, pdf, video, package]:
        require(path.is_file() and path.stat().st_size > 0, f"missing {path.name}")

    require(manifest["schema"] == "claimgate.osscontest-submission/v1", "unexpected manifest schema")
    require(manifest["reportBodyPageLimit"] == 5, "official report body limit must be five pages")
    require(package.stat().st_size == manifest["uploadZipBytes"], "manifest ZIP byte count mismatch")
    require(digest(package) == manifest["uploadZipSha256"], "manifest ZIP SHA-256 mismatch")
    for relative in manifest["licenseEvidence"]:
        require((OUT / relative).resolve().is_file(), f"missing license evidence: {relative}")
    for entry in manifest["files"]:
        path = OUT / entry["path"]
        require(path.stat().st_size == entry["bytes"], f"manifest byte count mismatch: {path.name}")
        require(digest(path) == entry["sha256"], f"manifest SHA-256 mismatch: {path.name}")

    report = Document(docx)
    text = "\n".join(p.text for p in report.paragraphs)
    text += "\n" + "\n".join(cell.text for table in report.tables for row in table.rows for cell in row.cells)
    require(manifest["repository"] in text, "report repository URL does not match manifest")
    require(manifest["videoUrl"] in text, "report video URL does not match manifest")
    for phrase in [
        "프로젝트 등록",
        "시연영상",
        "SBOM",
        "AI 모델 활용",
        "No Anchor",
        "Evidence Pack",
        "프로젝트 구조 및 코드 완성도",
        "오픈소스 프로젝트로의 발전 가능성",
        "개발 문서의 구체성",
        "프로젝트 혁신성",
        "프로젝트 협업 및 관리체계",
        "커뮤니티 확장 가능성",
        "오픈소스SW 적절성",
        "기능테스트",
        "라이선스 검증",
        "유형 2",
        "Apache-2.0",
        "productionQuality=false",
    ]:
        require(phrase.lower() in text.lower(), f"report phrase missing: {phrase}")

    page_text = pdf_pages(pdf)
    appendix_pages = [index for index, value in enumerate(page_text, 1) if "붙임 1" in value]
    require(appendix_pages, "could not find the SBOM appendix page")
    body_pages = appendix_pages[0] - 1
    require(1 <= body_pages <= manifest["reportBodyPageLimit"], f"report body is {body_pages} pages")
    require(any("붙임 2" in value and "AI 모델" in value for value in page_text), "AI model appendix missing")

    probe = json.loads(subprocess.check_output([
        "ffprobe", "-v", "error", "-show_entries", "stream=width,height", "-show_entries", "format=duration", "-of", "json", str(video)
    ], text=True))
    duration = float(probe["format"]["duration"])
    require(1 <= duration <= 180, f"video duration {duration}s exceeds 3 minutes")
    stream = probe["streams"][0]
    require(stream["width"] >= 1280 and stream["height"] >= 720, "video resolution below 1280x720")

    with ZipFile(package) as archive:
        require(archive.testzip() is None, "ZIP CRC failure")
        expected = [entry["path"] for entry in manifest["files"]]
        require(sorted(archive.namelist()) == sorted(expected), "official upload ZIP must contain the manifest DOCX and PDF only")
        for entry in manifest["files"]:
            require(hashlib.sha256(archive.read(entry["path"])).hexdigest() == entry["sha256"], f"ZIP content mismatch: {entry['path']}")

    blockers = []
    receipt = str(manifest["receiptNumber"]).strip()
    if not receipt or receipt == "접수번호" or "REQUIRED" in receipt:
        blockers.append("RECEIPT_NUMBER_REQUIRED")
    video_url = str(manifest["videoUrl"]).strip()
    if not re.fullmatch(r"https://(?:www\.)?(?:youtube\.com/watch\?v=[A-Za-z0-9_-]+|youtu\.be/[A-Za-z0-9_-]+)(?:[?&][^\s]+)?", video_url):
        blockers.append("YOUTUBE_URL_REQUIRED")
    repository_blocker = public_repository_blocker(str(manifest["repository"]).strip())
    if repository_blocker:
        blockers.append(repository_blocker)
    try:
        ahead = int(subprocess.check_output(["git", "rev-list", "--count", "@{u}..HEAD"], text=True).strip())
    except subprocess.CalledProcessError:
        blockers.append("PUBLIC_REPOSITORY_UPSTREAM_MISSING")
    else:
        if ahead:
            blockers.append(f"PUBLIC_REPOSITORY_PUSH_{ahead}_COMMITS")

    print(
        f"osscontest-final: ARTIFACT PASS — {body_pages} body pages + "
        f"{len(page_text) - body_pages} annex pages, {duration:.2f}s video, verified 2-file ZIP"
    )
    if blockers:
        print(f"osscontest-final: READINESS BLOCKED ({len(blockers)}) — " + ", ".join(blockers))
    else:
        print("osscontest-final: READY — receipt, YouTube URL, public repository, and pushed HEAD verified")


if __name__ == "__main__":
    main()
